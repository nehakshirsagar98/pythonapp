from docx.api import Document
import io
import string
import re
# Load the first table from your document. In your example file,
# there is only one table, so I just grab the first one.
filename = "D:\\work\\project\\ResumeClassification\\rc\\resume\\neha.docx"
other_text="this is example other text  vikram"
def get_tables(file_name,other_text):
    main_txt = []
    complete_str=""
    document = Document(file_name)
    print(len(document.tables))
    textFilename = file_name + ".txt"
    with io.open(textFilename, "w", encoding="utf-8") as textFile:
        for index in range(len(document.tables)):
            print(index)
            table = document.tables[index]

            # Data will be a list of rows represented as dictionaries
            # containing each row's data.
            data = []

            keys = None
            for i, row in enumerate(table.rows):
                text = (cell.text.strip() for cell in row.cells)
                # Establish the mapping based on the first row
                # headers; these will become the keys of our dictionary
                if i == 0:
                    keys = tuple(text)
                    continue

                # Construct a dictionary for this row, mapping
                # keys to values for this row
                row_data = dict(zip(keys, text))
                data.append(row_data)
            
            complete_str=complete_str+str(data)
        textFile.write(" " + other_text)
        #complete_str = complete_str.replace('\n',' ').replace('\r', ' ')
        complete_str = re.sub('[^a-zA-Z0-9+ \.]', ' ', complete_str)
        try:
            #print(' '.join(keys))
            #print(' '.join(text)) 
            #textFile.write(str(zip(keys, text)))    
            textFile.write(complete_str)
            
            #textFile.write(' '.join(text)+' ')    
        except KeyError as e:
            print(e)
            print(complete_str)
    return textFilename

print(get_tables(filename,other_text))